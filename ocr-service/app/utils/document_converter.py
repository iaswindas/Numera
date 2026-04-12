"""Convert Word (.docx) and Excel (.xlsx) files to structured text."""

from __future__ import annotations

import io
import logging
from typing import NamedTuple

logger = logging.getLogger("ocr-service.utils.document_converter")


class PageText(NamedTuple):
    page: int
    text: str


def is_docx(data: bytes) -> bool:
    """Check if data is a DOCX file (ZIP with word/ directory)."""
    return data[:4] == b"PK\x03\x04" and b"word/" in data[:2000]


def is_xlsx(data: bytes) -> bool:
    """Check if data is an XLSX file (ZIP with xl/ directory)."""
    return data[:4] == b"PK\x03\x04" and b"xl/" in data[:2000]


def extract_docx_text(data: bytes) -> list[PageText]:
    """Extract text from a DOCX file.

    Returns one PageText per paragraph group (treated as a single "page"
    since DOCX doesn't have physical page breaks in the XML model).
    """
    from docx import Document

    doc = Document(io.BytesIO(data))
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("\t".join(cells))
        if rows:
            paragraphs.append("\n".join(rows))

    if not paragraphs:
        return [PageText(page=1, text="")]

    return [PageText(page=1, text="\n".join(paragraphs))]


def extract_xlsx_text(data: bytes) -> list[PageText]:
    """Extract text from an XLSX file.

    Each worksheet becomes one "page". Cell values are tab-separated per row.
    """
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    pages: list[PageText] = []

    for idx, ws in enumerate(wb.worksheets, start=1):
        rows: list[str] = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(v) if v is not None else "" for v in row]
            if any(c for c in cells):
                rows.append("\t".join(cells))
        pages.append(PageText(page=idx, text="\n".join(rows)))

    wb.close()
    return pages if pages else [PageText(page=1, text="")]
